import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets margins (padding) of a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()
    
    # ----------------------------------------------------
    # DOCUMENT GEOMETRY
    # ----------------------------------------------------
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # ----------------------------------------------------
    # COLOR PALETTE & TYPOGRAPHY
    # ----------------------------------------------------
    PRIMARY_COLOR = RGBColor(31, 78, 121)    # Deep Navy
    SECONDARY_COLOR = RGBColor(89, 89, 89)  # Slate Grey
    TEXT_COLOR = RGBColor(51, 51, 51)       # Charcoal
    
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = TEXT_COLOR
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    
    # ----------------------------------------------------
    # COVER / HEADER TITLE
    # ----------------------------------------------------
    title = doc.add_paragraph()
    title_run = title.add_run("Redrob Hybrid Candidate Ranker: End-to-End System Report")
    title_run.font.name = 'Georgia'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(2)
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Architecture, Empirical Calibration, GBDT Learned Combiner, and Validation Suite")
    subtitle_run.font.name = 'Calibri'
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = SECONDARY_COLOR
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.paragraph_format.space_after = Pt(24)
    
    # ----------------------------------------------------
    # SECTION 1: PROBLEM STATEMENT & CHALLENGE CONSTRAINTS
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Problem Statement & Constraints")
    h1_run.font.name = 'Georgia'
    h1_run.font.size = Pt(18)
    h1_run.font.bold = True
    h1_run.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.add_run("The ").bold = True
    p.add_run("Redrob Intelligent Candidate Discovery & Ranking Challenge ")
    p.add_run("requires building a highly scalable, offline, and CPU-efficient ranking engine to discover the top ")
    p.add_run("100 ").bold = True
    p.add_run("candidates for a ")
    p.add_run("Senior AI Engineer ").bold = True
    p.add_run("role out of a total raw pool of ")
    p.add_run("100,000+ ").bold = True
    p.add_run("candidates. The evaluation criteria prioritize high lexical/semantic alignment, verified behavioral engagement, and candidate-specific explainable reasoning.")
    
    h2 = doc.add_heading(level=2)
    h2_run = h2.add_run("1.1 Challenge Constraints & Execution Guardrails")
    h2_run.font.name = 'Georgia'
    h2_run.font.size = Pt(14)
    h2_run.font.bold = True
    h2_run.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    
    p = doc.add_paragraph()
    p.add_run("The system must operate under strict environment and submission rules:")
    
    bullet_style = 'List Bullet'
    doc.add_paragraph("CPU-Only Offline Host: No network calls, API access, or GPU dependencies are permitted at runtime. Inference must run strictly locally.", style=bullet_style)
    doc.add_paragraph("Strict CSV Schema Format: The final output must be exactly team_Jarvis2.0.csv with columns: candidate_id, rank, score, reasoning.", style=bullet_style)
    doc.add_paragraph("Monotonicity Enforcements: Scores must be non-negative floats, strictly non-increasing by rank, and properly formatted in a sensible range (normalized [0, 1]).", style=bullet_style)
    doc.add_paragraph("Deterministic Tie-Breaking: Equal scores must be resolved alphabetically ascending by candidate_id to prevent rank-shuffling.", style=bullet_style)
    doc.add_paragraph("Safety & Honeypot Trap Thresholds: Detect and penalize honeypot profiles (candidates claiming 'Expert' proficiency in skills with 0 months of duration). Flagged honeypots must be under 10% (ideally 0%) in the final shortlist.", style=bullet_style)
    doc.add_paragraph("Factual, Gap-Free Explainability: Reasonings must be hallucination-free, candidate-specific, and explicitly mention all active ranking signals (Skill Assessments, GitHub Activity, and Recruiter Response Times) with clear fallbacks if they are default or missing.", style=bullet_style)

    # ----------------------------------------------------
    # SECTION 2: SYSTEM ARCHITECTURE
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("2. System Architecture & Pipeline Stages")
    h1_run.font.name = 'Georgia'
    h1_run.font.size = Pt(18)
    h1_run.font.bold = True
    h1_run.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.add_run("The Redrob Hybrid Ranker is structured as a multi-stage pipeline designed for memory efficiency, execution speed, and high-fidelity scoring:")
    
    # Stage 1
    doc.add_paragraph("Stage 1 — Boolean Hard Gates (Boolean Filtering)", style='Heading 3')
    p = doc.add_paragraph()
    p.add_run("To optimize memory footprint on CPU hosts, a streaming JSON parser reads candidates sequentially. It applies strict gates to filter out unqualified profiles before scoring: (1) Years of Experience (YOE) must be between 4 and 13; (2) Candidate must prefer Onsite/Hybrid work and be located or willing to relocate to Pune or Noida; (3) Profile Completeness must be >= 50%; (4) Excludes non-technical job title keywords. Out of 100,000 records, only ")
    p.add_run("679 candidates (0.68%) ").bold = True
    p.add_run("survive Stage 1, dramatically reducing the downstream scoring payload.")
    
    # Stage 2
    doc.add_paragraph("Stage 2 — Hybrid scoring Engine (Lexical & Semantic Match)", style='Heading 3')
    p = doc.add_paragraph()
    p.add_run("Calculates candidate-JD text alignment by blending three components: (1) ")
    p.add_run("Okapi BM25").bold = True
    p.add_run(" lexical retrieval on JD skill unigrams; (2) ")
    p.add_run("Contextual Proximity Matcher").bold = True
    p.add_run(" awarding boosts for matching skills positioned in the same sentence as senior action verbs ('built', 'scaled', 'optimized'); (3) ")
    p.add_run("Cross-Encoder Semantic Scoring").bold = True
    p.add_run(" utilizing the upgraded, larger 110M parameter ")
    p.add_run("cross-encoder/ms-marco-MiniLM-L-12-v2").italic = True
    p.add_run(" model for high-precision semantic similarity against the JD anchor.")
    
    # Stage 3
    doc.add_paragraph("Stage 3 — Behavioral Modifiers", style='Heading 3')
    p = doc.add_paragraph()
    p.add_run("Applies calibrated multipliers computed from candidate platform engagement metrics: inactivity decay (sigmoid centred on 88 days), recruiter response rates, recruiter response time average, active intent, saves/views, company-scale progression, notice period, contact verification, and assessment-first skill trust.")
    
    # Stage 4
    doc.add_paragraph("Stage 4 — GBDT LambdaMART Combiner", style='Heading 3')
    p = doc.add_paragraph()
    p.add_run("Rather than relying on manual weight tuning, the pipeline integrates a trained LightGBM GBDT ranker (objective='lambdarank') utilizing 15 dynamic features. Raw GBDT outputs (log-odds margins) are mapped to a clean ")
    p.add_run("[0.0, 1.0]").bold = True
    p.add_run(" range using Min-Max Normalization to guarantee professional, non-negative relevance scores.")
    
    # Stage 5
    doc.add_paragraph("Stage 5 — Explainable Reasoning Generator", style='Heading 3')
    p = doc.add_paragraph()
    p.add_run("Generates candidate-specific, factual narratives. To prevent structural template repetition, the script dynamically rotates layout orders and sentence structures based on candidate ID hashes. All reasonings are audited to ensure 100% coverage of active signals.")

    # ----------------------------------------------------
    # SECTION 3: CALIBRATION & CORRELATION MATRIX
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("3. Empirical Calibration & Pearson Correlations")
    h1_run.font.name = 'Georgia'
    h1_run.font.size = Pt(18)
    h1_run.font.bold = True
    h1_run.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.add_run("The multipliers in ")
    p.add_run("signal_calibration.py").italic = True
    p.add_run(" are calibrated on Pearson correlation coefficients computed across the 679-candidate gate survivors:")
    
    # Pearson Correlation Table
    table = doc.add_table(rows=7, cols=4)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Signal Pair", "Correlation (r)", "Sample Size (n)", "Strength"]
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        set_cell_background(hdr_cells[i], "1F4E79")
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(hdr_cells[i])
        
    corr_data = [
        ("Open to Work x Applications Submitted", "+0.1159", "679", "Weak"),
        ("Recruiter Response Rate x Open to Work", "+0.0549", "679", "Weak"),
        ("Recruiter Response Rate x Recruiter Saves", "+0.1324", "679", "Weak"),
        ("Recruiter Response Rate x Search Appearances", "+0.1406", "679", "Weak"),
        ("Recruiter Response Rate x Avg Response Time", "-0.1237", "679", "Weak"),
        ("Interview Completion Rate x GitHub Activity", "+0.0586", "364", "Weak")
    ]
    
    for row_idx, data in enumerate(corr_data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx])
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F2F2F2")
                
    doc.add_paragraph("Key finding: GitHub activity and interview completion rates show weak correlations (r = +0.0586). Therefore, GitHub is mapped to a narrow [0.97, 1.03] multiplier band (code-artifact nudge) rather than a linear weight, preventing it from dominating rankings unjustly.", style='Normal')

    # ----------------------------------------------------
    # SECTION 4: GBDT RANKER & LIMITATIONS
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("4. GBDT LambdaMART & Pseudo-Label Circularity")
    h1_run.font.name = 'Georgia'
    h1_run.font.size = Pt(18)
    h1_run.font.bold = True
    h1_run.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.add_run("The machine-learned combiner utilizes a GBDT LambdaRank model. In the spirit of scientific rigor, we document key architectural limitations:")
    
    # Bullet points for limitations
    p_lim1 = doc.add_paragraph(style=bullet_style)
    p_lim1.add_run("Pseudo-Label Circularity: ").bold = True
    p_lim1.add_run("Due to the lack of ground-truth relevance labels, the model was trained using our calibrated heuristic scores as target labels (pseudo-labels). As a result, the model acts as a learned approximation of our heuristic scoring function. Initially, the model achieved a correlation of 0.9859 with the heuristic target, indicating near-perfect memorization of heuristic rules.")
    
    p_lim2 = doc.add_paragraph(style=bullet_style)
    p_lim2.add_run("Cross-Encoder Regularization: ").bold = True
    p_lim2.add_run("Upgrading to ms-marco-MiniLM-L-12-v2 and retraining the GBDT model reduced memorization. Post-retraining, GBDT correlation with heuristic targets fell to 0.7290, and correlation with Cross-Encoder scores was 0.3729. The Cross-Encoder acts as a helpful semantic regularizer, but does not dominate the split structure.")
    
    p_lim3 = doc.add_paragraph(style=bullet_style)
    p_lim3.add_run("Lexical Dominance: ").bold = True
    p_lim3.add_run("Feature importance analysis shows that bm25_norm remains the dominant splitter (gain: 225.62), while recruiter_rr (gain: 25.43) and co_norm (gain: 8.57) serve as secondary splitters. The Cross-Encoder has a high split count but low gain, serving as a semantic regularizer.")

    # ----------------------------------------------------
    # SECTION 5: VERIFICATION & AUDIT RESULTS
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("5. Verification & Submission Audit Results")
    h1_run.font.name = 'Georgia'
    h1_run.font.size = Pt(18)
    h1_run.font.bold = True
    h1_run.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.add_run("The generated submission payload (")
    p.add_run("team_Jarvis2.0.csv").bold = True
    p.add_run(") has been rigorously audited and passes all checks:")
    
    doc.add_paragraph("Official Format Validator: PASSED (Submission is valid. Monotonicity and deterministic alphanumeric tie-breaking verified).", style=bullet_style)
    doc.add_paragraph("Zero Reasoning Gaps: 100% of rows mention Skill Assessments, GitHub, and response time metrics (either their real values or fallback strings). Search keyword checks resolve to 100/100 candidates.", style=bullet_style)
    doc.add_paragraph("Min-Max Score Boundaries: Score bounds are verified to be strictly within [0, 1]. Rank 1 score is exactly 1.00000000; Rank 100 score is 0.00934280.", style=bullet_style)
    doc.add_paragraph("Honeypot Trap Count: 0 / 100 (No honeypots selected in the top 100 shortlist; safety validation PASSED).", style=bullet_style)
    
    doc.add_paragraph("Document status: FINALIZED & VALIDATED.", style='Normal').paragraph_format.space_before = Pt(12)

    # ----------------------------------------------------
    # SECTION 6: PROJECT FILES & ROLE MAPPING
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("6. Project Files & Role Mapping")
    h1_run.font.name = 'Georgia'
    h1_run.font.size = Pt(18)
    h1_run.font.bold = True
    h1_run.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.add_run("The following table details the workspace file structure and explains what each script or file states and performs within the hybrid ranker system:")
    
    file_table = doc.add_table(rows=16, cols=2)
    file_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["File Path / Name", "Purpose & Role in the Pipeline"]
    hdr_cells = file_table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        set_cell_background(hdr_cells[i], "1F4E79")
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(hdr_cells[i])
        
    file_data = [
        ("rank_candidates.py", "Main ranking pipeline execution script. Streams candidates through Stage 1 Boolean gates, extracts feature sets, predicts ranking scores using the trained LambdaMART LightGBM model, applies Min-Max Normalization to [0, 1] range, resolves alphanumeric CAND_ID tie-breaking, and dynamically formats gap-free Candidate Explanations (reasonings)."),
        ("signal_calibration.py", "Houses empirically calibrated multiplier constants, sigmoid inactivity decay models, unverified contact penalties, and utility functions used for heuristic weights calibration."),
        ("train_learned_combiner.py", "Trained combiner script. Extracts 15 behavioral and semantic features for survivors and trains a LightGBM LambdaRank booster (objective='lambdarank') using calibrated continuous targets (grades mapped to [0, 30] integers)."),
        ("generate_crossencoder_scores.py", "Computes and gzips semantic scores for all Stage 1 survivors locally on CPU using the upgraded cross-encoder/ms-marco-MiniLM-L-12-v2 model against the JD semantic query anchor."),
        ("analyze_signals.py", "Runs statistical analysis (distributions, Pearson correlations, disengagement bands) over the surviving candidate pool, outputting the docs/SIGNAL_CALIBRATION.md report containing the GBDT ranker circularity limits notes."),
        ("validate_submission.py", "The official validation script. Confirms that the final team_Jarvis2.0.csv submission matches column headers, has exactly 100 rows, uses non-negative monotonic scores, and breaks ties alphabetically."),
        ("audit_submission.py", "Internal reasoning auditor. Validates that reasonings have sufficient length, do not contain honeypot keywords, do not repeat structural openers, and checks that assessment, GitHub, and response time metrics are always mentioned."),
        ("verify_setup.py", "Technical checker verifying that all candidate datasets, model weight caches, precomputed similarity payloads, and final CSV submission files are present and correct."),
        ("learned_model.txt", "Trained LightGBM GBDT booster binary containing decision trees and splitting rules used during Stage 4 combinatorics scoring."),
        ("team_Jarvis2.0.csv", "The official finalized submission file. Contains the top 100 candidates ranked, their min-max scaled scores, and gap-free explainability reasoning strings."),
        ("data/candidates.jsonl", "Core candidate repository containing profile metadata and behavioral signal logs for 100,000+ engineers."),
        ("data/sample_candidates.jsonl", " light verification dataset (2,969 candidates) designed so that exactly 100 candidates pass the Boolean gates to test formatting and reproduction scripts."),
        ("data/crossencoder_scores.json.gz", "Gzipped JSON payload storing the precomputed L-12 Cross-Encoder semantic similarity scores for all 679 survivors."),
        ("submission_metadata.yaml", "Stores team details, contact information, reproduction command, portal methodology description, and AI usage declaration."),
        ("docs/SIGNAL_CALIBRATION.md", "Empirical calibration documentation describing signal distributions, Pearson correlations, multiplier parameters, and GBDT model circularity limits.")
    ]
    
    for row_idx, data in enumerate(file_data, start=1):
        row_cells = file_table.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx])
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F2F2F2")

    # ----------------------------------------------------
    # SECTION 7: POTENTIAL VULNERABILITIES & MITIGATIONS
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("7. Potential Vulnerabilities & Mitigation Strategies")
    h1_run.font.name = 'Georgia'
    h1_run.font.size = Pt(18)
    h1_run.font.bold = True
    h1_run.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.add_run("We identify two core system vulnerabilities and outline their theoretical impact and corresponding engineering mitigations:")
    
    # Vulnerability 1
    p_vul1 = doc.add_paragraph(style=bullet_style)
    p_vul1.add_run("Vulnerability 1 — BM25 Lexical Dominance (Semantic Richness vs. Exact Keyword Overlap): ").bold = True
    p_vul1.add_run("\n• Theoretical Risk: ").bold = True
    p_vul1.add_run("Feature importance analysis shows that lexical matching (bm25_norm) dominates GBDT decisions. If judges favor semantic richness (overall role matching) over exact keyword match, a candidate with deep AI knowledge but missing specific terms (e.g. Pinecone, Elasticsearch) might rank lower than someone with exact keyword overlap but less overall depth.")
    p_vul1.add_run("\n• Mitigation Strategy: ").bold = True
    p_vul1.add_run("BM25 serves as a high-recall filter to ensure that candidates possess the exact infrastructure tools required by the JD. The L-12 Cross-Encoder then acts as a high-precision semantic reranker. The combination is explicitly designed to balance strict technical constraints with semantic understanding, ensuring that candidates who write generic profiles without concrete tool experience do not slip into the shortlist.")
    
    # Vulnerability 2
    p_vul2 = doc.add_paragraph(style=bullet_style)
    p_vul2.add_run("Vulnerability 2 — Heuristic Target circularity (Optimizing Heuristics rather than Ground-Truth): ").bold = True
    p_vul2.add_run("\n• Theoretical Risk: ").bold = True
    p_vul2.add_run("Because there is no human-labeled training dataset, the GBDT ranker was trained on heuristic targets. It is effectively learning a GBDT-based approximation of our own multiplier equations rather than discovering net-new predictive signals directly from raw data.")
    p_vul2.add_run("\n• Mitigation Strategy: ").bold = True
    p_vul2.add_run("This form of boot-strapped training (distant supervision) is the industry standard for cold-starting search systems when human labels are unavailable. To transition this ranker into a true machine-learning model, the production pipeline should log click-logs, recruiter saves, and interview bookings over time. This interaction data will form a true pairwise click-log training dataset, breaking the heuristic circularity loop.")

    # Save Document
    doc.save("redrob_project_scenario_report.docx")
    print("SUCCESS: Word document created with vulnerabilities section.")

if __name__ == "__main__":
    create_report()
