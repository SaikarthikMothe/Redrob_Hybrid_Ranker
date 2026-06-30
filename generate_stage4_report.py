import docx
from collections import defaultdict
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

PRIMARY = RGBColor(31, 78, 121)    # Deep Navy
ACCENT = RGBColor(0, 112, 192)     # Bright Blue
SECONDARY = RGBColor(89, 89, 89)   # Slate Grey
TEXT = RGBColor(51, 51, 51)        # Charcoal
BG_HEADER = "1F4E79"               # Table header fill (navy)
BG_LIGHT = "EBF3FB"                # Light blue row fill

FEATURE_SPECS = [
    ("bm25_norm", "Stage 2 normalized lexical relevance"),
    ("crossencoder_score", "Stage 2 deep semantic similarity"),
    ("co_norm", "Stage 2 contextual co-occurrence"),
    ("skill_trust", "Assessment-backed skill evidence"),
    ("activity_decay", "Sigmoid recency/activity signal"),
    ("recruiter_rr", "Raw recruiter response-rate signal"),
    ("rt_norm", "Normalized inverse response time"),
    ("intent_score", "Open-to-work and applications"),
    ("market_validation", "Views, recruiter saves, appearances"),
    ("company_scale", "Employer scale and trajectory"),
    ("icr", "Interview completion rate"),
    ("notice_norm", "Notice-period operational band"),
    ("github_norm", "Normalized GitHub activity"),
    ("contact_verified", "Email/phone verification status"),
    ("oar", "Offer acceptance reliability"),
]


def model_gain_importance(model_path="learned_model.txt"):
    """Parse LightGBM's text model without requiring the lightgbm package."""
    gains = defaultdict(float)
    with open(model_path, encoding="utf-8") as model:
        lines = iter(model)
        for line in lines:
            if not line.startswith("split_feature="):
                continue
            feature_indices = [int(value) for value in line.split("=", 1)[1].split()]
            gain_line = next(lines)
            if not gain_line.startswith("split_gain="):
                raise ValueError("Unexpected LightGBM model layout after split_feature")
            split_gains = [float(value) for value in gain_line.split("=", 1)[1].split()]
            for feature_index, gain in zip(feature_indices, split_gains):
                gains[feature_index] += gain

    total_gain = sum(gains.values())
    return sorted(
        [
            (index, name, (100.0 * gains[index] / total_gain if total_gain else 0.0), description)
            for index, (name, description) in enumerate(FEATURE_SPECS)
        ],
        key=lambda item: item[2],
        reverse=True,
    )

def shd(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd_el = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd_el)

def cell_margins(cell, top=80, bottom=80, left=130, right=130):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, val in [('w:top', top), ('w:bottom', bottom),
                      ('w:left', left), ('w:right', right)]:
        nd = OxmlElement(name)
        nd.set(qn('w:w'), str(val))
        nd.set(qn('w:type'), 'dxa')
        tcMar.append(nd)
    tcPr.append(tcMar)

def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.name = 'Georgia'
    r.font.size = Pt(17)
    r.font.bold = True
    r.font.color.rgb = PRIMARY
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    return p

def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p

def body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_parts:
        remaining = text
        for bp in sorted(bold_parts, key=lambda x: text.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                run = p.add_run(remaining[:idx])
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = TEXT
            br = p.add_run(bp)
            br.font.name = 'Calibri'
            br.font.size = Pt(11)
            br.font.bold = True
            br.font.color.rgb = TEXT
            remaining = remaining[idx + len(bp):]
        if remaining:
            run = p.add_run(remaining)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = TEXT
    else:
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
        r.font.color.rgb = TEXT
    return p

def bullet(doc, text, level=0, color=TEXT):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.color.rgb = color
    return p

def main():
    doc = docx.Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Document Title Block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("Stage 4 Technical Report: GBDT Learned Combiner & Score Normalization")
    r_title.font.name = 'Georgia'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = PRIMARY
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sub.paragraph_format.space_after = Pt(20)
    r_sub = p_sub.add_run("Redrob Hybrid Candidate Ranker - Machine Learning Ranking Engine & Normalization Safeguards")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = SECONDARY
    
    # Section 1
    h1(doc, "1. Executive Summary")
    body(doc, "Stage 4 of the Redrob Hybrid Candidate Ranker implements a machine learning-based ranking combiner "
              "using a Gradient Boosted Decision Tree (GBDT) regression framework. Rather than relying on static heuristic "
              "feature weights, Stage 4 extracts a comprehensive 15-dimensional feature matrix for each surviving "
              "candidate profile and feeds it into a trained LightGBM booster. To guarantee that the output scores "
              "comply with strict competition guardrails and avoid grading invalidation, we have established a "
              "double-normalized and clipped score boundary scheme that restricts all outputs strictly within the "
              "[0.0, 1.0] range.",
              ["Stage 4", "15-dimensional feature matrix", "LightGBM booster", "[0.0, 1.0] range"])
              
    # Section 2
    h1(doc, "2. GBDT Feature Design & Model Architecture")
    body(doc, "The LightGBM ranking model integrates 15 distinct candidate features spanning lexical match, "
              "semantic similarity, contextual proximity, and platform engagement history. This matrix is not a "
              "simple concatenation of the Stage 2 signals and the Stage 3 report's nine heuristic multiplier groups. "
              "It is a learned representation composed of 3 relevance inputs and 12 behavioral/auxiliary inputs:")
              
    h2(doc, "2.1 15-Dimensional Feature Matrix")
    bullet(doc, "bm25_norm: Normalized lexical similarity score computed over the entire candidate pool.")
    bullet(doc, "crossencoder_score: Deep semantic similarity score against the JD anchor.")
    bullet(doc, "co_norm: Contextual co-occurrence score representing active execution evidence.")
    bullet(doc, "skill_trust: Assessment-first skill evidence confidence score.")
    bullet(doc, "activity_decay: Sigmoid-based recruiter platform activity decay score.")
    bullet(doc, "recruiter_rr: Recruiter response rate.")
    bullet(doc, "rt_norm: Normalized average recruiter response time.")
    bullet(doc, "intent_score: Compound job-seeking intent proxy (open to work, applications submitted).")
    bullet(doc, "market_validation: Recruiter interest metric (saves, views, appearances).")
    bullet(doc, "company_scale: Employer scale and growth trajectory progression delta.")
    bullet(doc, "icr: Interview completion rate.")
    bullet(doc, "notice_norm: Notice period contractual constraint rating.")
    bullet(doc, "github_norm: Public coding and open-source contribution score.")
    bullet(doc, "contact_verified: Dual-channel email/phone verification status.")
    bullet(doc, "oar: Historical offer acceptance rate.")

    h2(doc, "2.2 Cross-Stage Dimensional Reconciliation")
    body(doc, "The dimensional accounting is exact: Stage 2 contributes bm25_norm, crossencoder_score, and co_norm. "
              "Stage 4 then supplies 12 separately addressable behavioral and auxiliary columns. The Stage 3 count of "
              "nine refers to heuristic multiplier groups used to generate the continuous training target; it does not "
              "define the learned model's input width. In particular, OAR and ICR are separate model columns, response "
              "time is retained as rt_norm, intent and market validation are retained as distinct features, and "
              "company_scale represents employer trajectory. Thus, 3 + 12 = 15 with no hidden or legacy dimensions.",
              ["3 + 12 = 15", "no hidden or legacy dimensions"])
    
    h2(doc, "2.3 LightGBM Model Configuration")
    body(doc, "The regression model is trained on continuous targets utilizing the LightGBM library with shallow, "
              "regularized decision trees to avoid overfitting:")
    bullet(doc, "Objective: Regression (fitting continuous score targets).")
    bullet(doc, "Shallow Tree Depth: limited to 16 leaves (num_leaves=16) to ensure generalization on unseen datasets.")
    bullet(doc, "Minimum Data in Leaf: 10 candidates (min_data_in_leaf=10).")
    bullet(doc, "Regularization: Early stopping configured for 50 rounds (stopping_rounds=50) on validation splits.")

    # Section 3
    h1(doc, "3. Score Bounds & Normalization Safeguards")
    body(doc, "Raw outputs from regression boosters are unbounded continuous predictions, "
              "which can easily exceed 1.0 or fall below 0.0. To prevent grading invalidation, we implement a multi-layered boundary defense:")
    
    h2(doc, "3.1 Multi-Layered Normalization Pipeline")
    bullet(doc, "First-Pass Normalization: Raw booster outputs are min-max normalized across all survival pool candidates.")
    bullet(doc, "Second-Pass Normalization: Once the top 100 candidates are shortlisted, they are normalized again to ensure that Rank 1 has a score of exactly 1.00000000.")
    bullet(doc, "Boundary Clipping: A strict failsafe clip max(0.0, min(1.0, score)) is applied during normalisation and immediately before writing to the CSV to protect against floating-point leaks.")
    
    h2(doc, "3.2 Model Feature Importances (Gain)")
    body(doc, "The table below reports every input column used by train_learned_combiner.py and rank_candidates.py, "
              "sorted by realized model gain. The feature names map directly to the exact array order listed in Section 2.1. "
              "Gain percentages were calculated directly from "
              "learned_model.txt and sum to 100% (subject to displayed rounding). A 0.00% gain means the feature was "
              "present in the 15-column matrix but the fitted trees did not select it for a split.")
    
    # Importance Table
    table = doc.add_table(rows=16, cols=4)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    column_widths = [0.45, 1.45, 1.10, 3.50]
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Rank"
    hdr_cells[1].text = "Feature Name"
    hdr_cells[2].text = "Importance (Gain %)"
    hdr_cells[3].text = "Source / Core Description"
    for cell in hdr_cells:
        shd(cell, BG_HEADER)
        cell.width = Inches(column_widths[hdr_cells.index(cell)])
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell_margins(cell)
        
    importance_data = model_gain_importance()
    
    for row_idx, (_, name, gain, source) in enumerate(importance_data, 1):
        row_cells = table.rows[row_idx].cells
        row_cells[0].text = str(row_idx)
        row_cells[1].text = name
        row_cells[2].text = "<0.01%" if 0 < gain < 0.01 else f"{gain:.2f}%"
        row_cells[3].text = source
        bg_color = BG_LIGHT if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell in enumerate(row_cells):
            shd(cell, bg_color)
            cell.width = Inches(column_widths[col_idx])
            cell_margins(cell)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            
    body(doc, "") # Spacer
    body(doc, "Conclusion: Stage 4 integrates GBDT model combinations with strict score boundary controls. "
              "This ensures the ranking engine balances complex feature interactions while producing "
              "valid, non-negative scores at or below 1.0.",
              ["GBDT model combinations", "non-negative scores", "at or below 1.0"])
              
    doc.save("stage4_report.docx")
    print("SUCCESS: stage4_report.docx generated.")

if __name__ == "__main__":
    main()
